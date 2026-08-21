from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)

for s in ["Normal","Title","Heading 1","Heading 2"]:
    doc.styles[s].font.name = "Aptos"
doc.styles["Normal"].font.size = Pt(10.5)
doc.styles["Title"].font.size = Pt(24)
doc.styles["Heading 1"].font.size = Pt(16)
doc.styles["Heading 2"].font.size = Pt(12.5)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NEUROSHIELD AI"); r.bold = True; r.font.size = Pt(25)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Documentation of Individual Contribution"); r.bold = True; r.font.size = Pt(16)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("An Explainable Machine Learning Framework for Stroke Risk Prediction").italic = True

table = doc.add_table(rows=4, cols=2); table.style = "Light Shading Accent 1"
for i,(a,b) in enumerate([
    ("Team Member","Pragati Sharma"),
    ("Role","Team Lead & ML Integration"),
    ("Primary Area","Machine Learning Model Development, Training & Integration"),
    ("Project","NeuroShield AI")]):
    table.cell(i,0).text=a; table.cell(i,1).text=b

doc.add_heading("1. Overview of My Contribution",1)
doc.add_paragraph(
"As the Team Lead and ML Integration member of NeuroShield AI, my primary responsibility was to "
"take the prepared dataset through the machine-learning model development stage, train and compare "
"the selected models, perform model selection and save the final trained model. I also worked on "
"integrating the work produced by the other team members and maintaining the project workflow through Git/GitHub."
)
doc.add_paragraph(
"The defined ML tasks for my role were Logistic Regression, Decision Tree, Random Forest, model "
"comparison, hyperparameter tuning and final model selection. My deliverables included the final "
"trained model, model-comparison results, an integrated notebook and the final presentation."
)

doc.add_heading("2. Machine Learning Model Development",1)
doc.add_heading("2.1 Models Evaluated",2)
for x in ["Logistic Regression","Decision Tree","Random Forest"]:
    doc.add_paragraph(x, style="List Bullet")
doc.add_paragraph(
"These models were used to establish and compare classical machine-learning approaches for stroke-risk prediction."
)
doc.add_heading("2.2 Model Training",2)
doc.add_paragraph(
"After preprocessing and exploratory analysis, I used the prepared data for model training. "
"The trained models were compared as part of the model-selection process."
)
doc.add_heading("2.3 Hyperparameter Tuning",2)
doc.add_paragraph(
"Hyperparameter tuning was included to improve the configuration of the selected model and establish "
"a stronger final baseline. Exact final parameter values should be taken from the final training notebook."
)

doc.add_heading("3. Final Model Artifact",1)
doc.add_paragraph("The project produced a serialized trained model:")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("models/stroke_model.pkl"); r.bold=True
doc.add_paragraph(
"The model artifact is intended for later integration into the prediction service. The project "
"documentation defines a future pathway in which the trained model can be exposed through a Python API "
"and consumed by the application backend."
)

doc.add_heading("4. Integration Responsibilities",1)
for x in [
"Project planning and GitHub repository management.",
"Coordination of the handoff between preprocessing, EDA, training and evaluation.",
"Integration of the team's notebooks and model artifacts.",
"Preparation of the final project presentation and demonstration.",
"Documentation and maintenance of the overall project workflow."
]: doc.add_paragraph(x, style="List Bullet")

doc.add_heading("5. End-to-End ML Workflow",1)
for i,x in enumerate([
"Data Collection & Preprocessing — Khushboo",
"Exploratory Data Analysis & Visualization — Nitya",
"Model Building & Training — Pragati",
"Model Evaluation & Explainable AI — Yashika",
"Project Integration, Documentation & Presentation — Pragati"],1):
    doc.add_paragraph(f"{i}. {x}")

doc.add_heading("6. Collaboration With Team Members",1)
doc.add_paragraph(
"Khushboo prepared the cleaned dataset and preprocessing workflow. Nitya worked on exploratory "
"analysis and visualizations. I used the processed data for model building and training. Yashika "
"handled model evaluation and Explainable AI using SHAP. I then combined the project components "
"for integration and documentation."
)

doc.add_heading("7. Repository Structure Relevant to My Work",1)
for x in [
"notebooks/03_model_training.ipynb — model training work",
"models/stroke_model.pkl — trained model artifact",
"reports/ — figures, results and paper material",
"requirements.txt — project dependencies",
"README.md — project documentation"
]: doc.add_paragraph(x, style="List Bullet")

doc.add_heading("8. Current Status",1)
table=doc.add_table(rows=1,cols=2); table.style="Light Shading Accent 1"
table.rows[0].cells[0].text="Area"; table.rows[0].cells[1].text="Status"
for a,b in [
("Model development","Completed"),
("Baseline ML training","Completed"),
("Model comparison","Completed as part of the project workflow"),
("Final trained model artifact","Created: stroke_model.pkl"),
("Git/GitHub integration","Completed for current ML work"),
("Project integration/documentation","Being consolidated"),
("Deep Learning advancement","Next-stage work"),
("Backend/frontend deployment","Next-stage work")]:
    c=table.add_row().cells; c[0].text=a; c[1].text=b

doc.add_heading("9. Research Contribution",1)
doc.add_paragraph(
"My contribution provides the core model-building and training stage of NeuroShield AI. It establishes "
"the predictive baseline that can subsequently be evaluated using appropriate metrics, interpreted using "
"Explainable AI techniques, and extended with Deep Learning. The trained model also provides the foundation "
"for later deployment as part of a full-stack healthcare decision-support platform."
)

doc.add_heading("10. Next Work Planned",1)
for x in [
"Deep Learning advancement and controlled comparison with the existing ML baseline.",
"Final consolidation of verified evaluation metrics for the research paper.",
"Backend/API integration for model inference.",
"Frontend integration and end-to-end testing.",
"Preparation of the final SIH 2026 prototype and demonstration."
]: doc.add_paragraph(x, style="List Bullet")

doc.add_heading("11. Reporting Note",1)
doc.add_paragraph(
"Numerical performance values such as accuracy, precision, recall, F1-score and ROC-AUC are not "
"included here because verified final values were not present in the available project documentation. "
"They should be inserted only from the final evaluation notebook/results."
)

doc.add_heading("12. Summary",1)
doc.add_paragraph(
"Pragati's contribution to NeuroShield AI centers on model building, training, comparison, final model "
"selection, project integration, version-control coordination and final documentation/presentation. "
"This completed ML foundation supports the next phases of Deep Learning, explainability refinement, "
"backend/frontend integration and SIH 2026 preparation."
)

path="/mnt/data/NeuroShield_AI_Pragati_Contribution_Documentation.docx"
doc.save(path)
print(path)
